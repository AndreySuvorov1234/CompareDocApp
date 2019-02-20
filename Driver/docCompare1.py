from Driver import difflib
import docx


class DocCompare:

    def __init__(self, doc1=False, doc2=False):
        self.docPath1 = doc1
        self.docPath2 = doc2
        self.text1 = ""
        self.text2 = ""
        self.text1_words = ""
        self.text2_words = ""
        self.comparison = float(0)

    def loadDocs(self, path1=False, path2=False):
        success = False
        if path1:
            self.docPath1 = path1
            success = True
        if path2:
            self.docPath2 = path2
            success = True
        return success

    def loadText(self, stext1=False, stext2=False):
        success = False
        if stext1:
            self.text1_words = stext1.split()
            success = True
        if stext2:
            self.text2_words = stext2.split()
            success = True
        return success


    def parseDocs(self):
        if self.docPath1 and self.docPath2:
            ext1 = self.docPath1.split(".")[-1]
            ext2 = self.docPath2.split(".")[-1]
            if (ext1 == "doc" or ext1 == "docx") and (ext2 == "doc" or ext2 == "docx"):
                self.openWordFiles()
            elif ext1 == "txt" and ext2 == "txt":
                self.openTextFiles()
            success = self.loadText(self.text1, self.text2)
            return success
        else:
            return False

    def openTextFiles(self):
        docA = open(self.docPath1, "r")
        self.text1 = docA.read()
        docB = open(self.docPath2, "r")
        self.text2 = docB.read()


    def openWordFiles(self):
        doc1 = docx.Document(self.docPath1)
        textList1 = [para.text for para in doc1.paragraphs]
        self.text1 = "\n".join(textList1)

        doc2 = docx.Document(self.docPath2)
        textList2 = [para.text for para in doc2.paragraphs]
        self.text2 = "\n".join(textList2)

    def compareDocs(self):
        if self.text1_words and self.text2_words:
            result = difflib.SequenceMatcher(None, self.text1_words, self.text2_words).ratio()
            self.comparison = result
            return True
        else:
            return False

    def run(self):
        self.parseDocs()
        self.compareDocs()

    def lastComparison(self):
        return self.comparison


#
# import time
# comparison = DocCompare("C:/Users/Andrey/PycharmProjects/CompareDocs/SampleComparisonDocs/Doc1.txt", "C:/Users/Andrey/PycharmProjects/CompareDocs/SampleComparisonDocs/Doc2.txt")
# comparison = DocCompare("C:/Users/Andrey/OneDrive/resume2.docx", "C:/Users/Andrey/OneDrive/resume.docx")
#
# print(comparison.docPath1)
# start = time.time()
# print(start)
# comparison.run()
# end = time.time()
# print(end-start)
# print(comparison.lastComparison())
# print(comparison.docPath1)
# print(comparison.docPath2)
# print(comparison.text1)
# print(comparison.text2)
# print(comparison.text1_words)
# print(comparison.text2_words)


